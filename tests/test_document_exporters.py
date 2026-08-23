"""Tests for document exporters, protocol conformance, and safe table traversal."""

from __future__ import annotations

import io

import pytest

from omniscribe.core.block_tree import (
    BlockNode,
    BlockType,
    DocumentTree,
    PageTree,
    TableNode,
)
from omniscribe.core.document import DocumentBlock, DocumentPage, DocumentResult
from omniscribe.core.document_exporters import (
    BaseDocumentExporter,
    DocumentExportProtocol,
)
from omniscribe.core.docx_tree_writer import DocxTreeExporter, convert_tree_to_docx
from omniscribe.core.docx_writer import DocxMarkdownExporter
from omniscribe.core.html_writer import HtmlExporter, render_html
from omniscribe.core.workflows.utils import validate_bbox_coordinates


def test_document_export_protocol_conformance() -> None:
    """Verify that all exporter implementations satisfy the runtime DocumentExportProtocol."""
    tree_exporter = DocxTreeExporter()
    md_exporter = DocxMarkdownExporter()
    html_exporter = HtmlExporter()

    assert isinstance(tree_exporter, DocumentExportProtocol)
    assert isinstance(md_exporter, DocumentExportProtocol)
    assert isinstance(html_exporter, DocumentExportProtocol)
    assert isinstance(tree_exporter, BaseDocumentExporter)
    assert isinstance(md_exporter, BaseDocumentExporter)
    assert isinstance(html_exporter, BaseDocumentExporter)


def test_docx_tree_exporter_polymorphic_export() -> None:
    """Verify DocxTreeExporter handles both DocumentTree and DocumentResult."""
    exporter = DocxTreeExporter()

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=(0.0, 0.0, 1.0, 0.5),
                        text="Sample text from tree",
                        page_idx=0,
                    )
                ],
            )
        ]
    )

    doc_result = DocumentResult(
        pages=[
            DocumentPage(
                page_index=0,
                blocks=[
                    DocumentBlock(
                        bbox=(0.0, 0.0, 1.0, 0.5),
                        text="Sample text from DocumentResult",
                    )
                ],
            )
        ]
    )

    # export_tree
    res_tree = exporter.export_tree(tree)
    assert isinstance(res_tree, io.BytesIO)
    assert len(res_tree.getvalue()) > 0

    # export_document
    res_doc = exporter.export_document(doc_result)
    assert isinstance(res_doc, io.BytesIO)
    assert len(res_doc.getvalue()) > 0

    # polymorphic export
    res_poly_tree = exporter.export(tree)
    assert isinstance(res_poly_tree, io.BytesIO)
    res_poly_doc = exporter.export(doc_result)
    assert isinstance(res_poly_doc, io.BytesIO)


def test_html_exporter_polymorphic_export() -> None:
    """Verify HtmlExporter handles both DocumentTree and DocumentResult."""
    exporter = HtmlExporter()

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.SECTION_HEADER,
                        bbox=(0.0, 0.0, 1.0, 0.2),
                        text="Title Heading",
                        page_idx=0,
                    )
                ],
            )
        ]
    )

    doc_result = DocumentResult(
        pages=[
            DocumentPage(
                page_index=0,
                blocks=[
                    DocumentBlock(
                        bbox=(0.0, 0.0, 1.0, 0.2),
                        text="Heading from doc",
                    )
                ],
            )
        ]
    )

    html_from_tree = exporter.export_tree(tree)
    assert "Title Heading" in html_from_tree
    assert "<!DOCTYPE html>" in html_from_tree

    html_from_doc = exporter.export_document(doc_result)
    assert "Heading from doc" in html_from_doc

    poly_tree = exporter.export(tree)
    assert "Title Heading" in poly_tree


def test_docx_markdown_exporter_polymorphic_export() -> None:
    """Verify DocxMarkdownExporter handles both DocumentTree and DocumentResult."""
    exporter = DocxMarkdownExporter()

    doc_result = DocumentResult(
        pages=[
            DocumentPage(
                page_index=0,
                blocks=[
                    DocumentBlock(
                        bbox=(0.0, 0.0, 1.0, 0.5),
                        text="# Heading\n\nParagraph text",
                    )
                ],
            )
        ]
    )

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=(0.0, 0.0, 1.0, 0.5),
                        text="Tree content",
                        page_idx=0,
                    )
                ],
            )
        ]
    )

    res_doc = exporter.export_document(doc_result)
    assert isinstance(res_doc, io.BytesIO)
    assert len(res_doc.getvalue()) > 0

    res_tree = exporter.export_tree(tree)
    assert isinstance(res_tree, io.BytesIO)


def test_safe_table_traversal_missing_cells_and_rows() -> None:
    """Verify docx and html exporters handle malformed or missing cells/rows without crashing."""
    # 1. BlockNode with BlockType.TABLE but no cells/rows attribute
    malformed_table_block = BlockNode(
        block_type=BlockType.TABLE,
        bbox=(0.0, 0.2, 1.0, 0.5),
        text="Fallback table text content",
        page_idx=0,
    )

    # 2. TableNode with empty cells list
    empty_table_node = TableNode(
        rows=0,
        cols=0,
        page_idx=0,
        bbox=(0.0, 0.5, 1.0, 0.8),
        cells=[],
    )

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[malformed_table_block, empty_table_node],
            )
        ]
    )

    # DOCX export should safely render without throwing AttributeError
    docx_stream = convert_tree_to_docx(tree)
    assert isinstance(docx_stream, io.BytesIO)
    assert len(docx_stream.getvalue()) > 0

    # HTML export should render safely
    html_out = render_html(tree)
    assert isinstance(html_out, str)
    assert "Fallback table text content" in html_out


def test_table_deduplication_in_docx_and_html() -> None:
    """Verify that a table present in both page.children and tree.tables is only rendered once."""
    table_node = TableNode(
        rows=1,
        cols=1,
        page_idx=0,
        bbox=(0.0, 0.1, 1.0, 0.3),
        cells=[
            [
                BlockNode(
                    block_type=BlockType.TEXT,
                    bbox=(0.0, 0.1, 1.0, 0.3),
                    text="Unique Table Cell Content",
                    page_idx=0,
                )
            ]
        ],
        block_id="unique_table_block_123",
    )

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[table_node],
            )
        ],
        tables=[table_node],  # Duplicate in tree.tables
    )

    from docx import Document

    docx_stream = convert_tree_to_docx(tree)
    doc = Document(docx_stream)
    # Must only have 1 table, not 2
    assert len(doc.tables) == 1

    html_out = render_html(tree)
    # Must only have 1 occurrence of the table in HTML
    assert html_out.count("<table") == 1
    assert "Unique Table Cell Content" in html_out


def test_validate_bbox_coordinates() -> None:
    """Verify bounding box validation and clamping logic."""
    # Valid box
    assert validate_bbox_coordinates([0.1, 0.2, 0.8, 0.9]) == (0.1, 0.2, 0.8, 0.9)

    # Clamping out-of-bounds
    assert validate_bbox_coordinates([-0.5, -0.2, 1.5, 2.0], clamp=True) == (
        0.0,
        0.0,
        1.0,
        1.0,
    )

    # Inverted coordinates ordered properly
    assert validate_bbox_coordinates([0.8, 0.9, 0.1, 0.2]) == (0.1, 0.2, 0.8, 0.9)

    # Invalid lengths or types raise ValueError
    with pytest.raises(ValueError, match="Expected 4 coordinate values"):
        validate_bbox_coordinates([0.1, 0.2, 0.3])

    with pytest.raises(ValueError, match="finite numbers"):
        validate_bbox_coordinates([0.1, float("nan"), 0.5, 0.8])
