"""Tests for :mod:`omniscribe.core.docx_tree_writer`."""

from __future__ import annotations

from omniscribe.core.block_tree import (
    BlockNode,
    BlockType,
    DocumentTree,
    PageTree,
    Span,
    TableNode,
)
from omniscribe.core.docx_tree_writer import convert_tree_to_docx


def test_docx_tree_writer_creates_real_structure():
    from docx import Document

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.SECTION_HEADER,
                        bbox=[0, 0, 1, 0.1],
                        text="Title",
                        page_idx=0,
                        level=1,
                    ),
                    BlockNode(
                        block_type=BlockType.LIST_ITEM,
                        bbox=[0, 0.1, 1, 0.2],
                        text="item",
                        page_idx=0,
                        level=0,
                    ),
                    BlockNode(
                        block_type=BlockType.CODE,
                        bbox=[0, 0.2, 1, 0.3],
                        text="x = 1",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.3, 1, 0.4],
                        text="body",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )
    stream = convert_tree_to_docx(tree)
    doc = Document(stream)
    # Heading 1 should be present
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    assert "Title" in headings
    # Code line preserved
    assert any("x = 1" in p.text for p in doc.paragraphs)
    # Body text present
    assert any("body" in p.text for p in doc.paragraphs)
    # List item present
    assert any("item" in p.text for p in doc.paragraphs)


def test_docx_tree_writer_renders_all_block_types():
    """Verify rendering of equations, key-value pairs, tables, rich text spans, and skipping headers/footers."""
    from docx import Document

    table_node = TableNode(
        rows=2,
        cols=2,
        page_idx=0,
        bbox=(0.0, 0.4, 1.0, 0.6),
        cells=[
            [
                BlockNode(
                    block_type=BlockType.TEXT,
                    bbox=(0.0, 0.4, 0.5, 0.5),
                    text="Header 1",
                    page_idx=0,
                ),
                BlockNode(
                    block_type=BlockType.TEXT,
                    bbox=(0.5, 0.4, 1.0, 0.5),
                    text="Header 2",
                    page_idx=0,
                ),
            ],
            [
                BlockNode(
                    block_type=BlockType.TEXT,
                    bbox=(0.0, 0.5, 0.5, 0.6),
                    text="Value 1",
                    page_idx=0,
                ),
                BlockNode(
                    block_type=BlockType.TEXT,
                    bbox=(0.5, 0.5, 1.0, 0.6),
                    text="Value 2",
                    page_idx=0,
                ),
            ],
        ],
    )

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    # 1. Equation
                    BlockNode(
                        block_type=BlockType.EQUATION,
                        bbox=(0.0, 0.0, 1.0, 0.1),
                        text="E=mc^2",
                        page_idx=0,
                    ),
                    # 2. Key-value pair
                    BlockNode(
                        block_type=BlockType.KEY_VALUE,
                        bbox=(0.0, 0.1, 1.0, 0.2),
                        text="Albert Einstein",
                        metadata={"key": "Author"},
                        page_idx=0,
                    ),
                    # 3. Rich text spans (bold, italic)
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=(0.0, 0.2, 1.0, 0.3),
                        text="",
                        spans=[
                            Span(text="Important: ", bold=True),
                            Span(text="relativity is "),
                            Span(text="universal", italic=True),
                        ],
                        page_idx=0,
                    ),
                    # 4. Table
                    table_node,
                    # 5. Header / Footer / Page number to skip
                    BlockNode(
                        block_type=BlockType.PAGE_HEADER,
                        bbox=(0.0, 0.0, 1.0, 0.05),
                        text="Confidential Header Should Be Skipped",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_FOOTER,
                        bbox=(0.0, 0.95, 1.0, 1.0),
                        text="Footer Note Should Be Skipped",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_NUMBER,
                        bbox=(0.9, 0.95, 1.0, 1.0),
                        text="Page 42 Should Be Skipped",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )

    stream = convert_tree_to_docx(tree)
    doc = Document(stream)

    # Verify Equation: rendered as '$ E=mc^2 $' with italic format
    equation_paragraphs = [p for p in doc.paragraphs if "$ E=mc^2 $" in p.text]
    assert len(equation_paragraphs) == 1
    assert equation_paragraphs[0].runs[0].italic is True

    # Verify Key-Value: rendered as 'Author: Albert Einstein' with bold key
    kv_paragraphs = [p for p in doc.paragraphs if "Author: Albert Einstein" in p.text]
    assert len(kv_paragraphs) == 1
    assert kv_paragraphs[0].runs[0].text == "Author: "
    assert kv_paragraphs[0].runs[0].bold is True
    assert kv_paragraphs[0].runs[1].text == "Albert Einstein"

    # Verify Rich text spans: bold and italic runs
    span_paragraphs = [
        p for p in doc.paragraphs if "Important: relativity is universal" in p.text
    ]
    assert len(span_paragraphs) == 1
    runs = span_paragraphs[0].runs
    assert runs[0].text == "Important: "
    assert runs[0].bold is True
    assert runs[1].text == "relativity is "
    assert runs[2].text == "universal"
    assert runs[2].italic is True

    # Verify Table: 2x2 table rendered with cells
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "Header 1"
    assert table.cell(0, 1).text == "Header 2"
    assert table.cell(1, 0).text == "Value 1"
    assert table.cell(1, 1).text == "Value 2"

    # Verify Skipping: page headers, footers, and page numbers are not in body text
    all_paragraph_text = " ".join(p.text for p in doc.paragraphs)
    assert "Confidential Header Should Be Skipped" not in all_paragraph_text
    assert "Footer Note Should Be Skipped" not in all_paragraph_text
    assert "Page 42 Should Be Skipped" not in all_paragraph_text
