"""Tests for the Phase 1-5 document-intelligence additions.

Covers the new modules added in the DocuVerse upgrade:

- :class:`omniscribe.core.block_tree` (DocumentTree IR)
- :class:`omniscribe.core.glossary` (Glossary)
- :class:`omniscribe.core.entity_memory` (EntityMemory)
- :class:`omniscribe.core.handwriting_preprocessor` (Sauvola + stroke + slant)
- :class:`omniscribe.core.translation_tree` (tree translation, sliding window)
- :class:`omniscribe.core.html_writer` (HTML export)
- :class:`omniscribe.core.docx_tree_writer` (DOCX export from tree)
- :class:`omniscribe.core.tree_export` (JSON export)
- :class:`omniscribe.core.dual_translator` (dual-translate)
- :class:`omniscribe.core.trocr_engine` (TrOCR lazy + heuristic confidence)
- :class:`omniscribe.core.nllb_engine` (NLLB language code resolver)
- :class:`omniscribe.core.translation_tree` (sliding window context)
"""

from __future__ import annotations

import asyncio

import numpy as np

from omniscribe.core.block_tree import (
    BlockNode,
    BlockType,
    DocumentTree,
    PageTree,
    TableNode,
    from_pages_data,
)
from omniscribe.core.docx_tree_writer import convert_tree_to_docx
from omniscribe.core.dual_translator import dual_translate
from omniscribe.core.entity_memory import EntityMemory
from omniscribe.core.glossary import Glossary, GlossaryEntry
from omniscribe.core.handwriting_preprocessor import (
    HandwritingOptions,
    estimate_stroke_width,
    is_handwritten_page,
    normalize_stroke_width,
    sauvola_binarize,
)
from omniscribe.core.html_writer import render_html
from omniscribe.core.nllb_engine import LANGUAGE_CODE_MAP, resolve_nllb_code
from omniscribe.core.translation_tree import (
    build_context_block,
    translate_tree,
)
from omniscribe.core.tree_export import export_json, export_json_bytes
from omniscribe.core.trocr_engine import _heuristic_confidence

# ---------------------------------------------------------------------------
# block_tree IR
# ---------------------------------------------------------------------------


def test_block_node_to_dict_round_trip():
    node = BlockNode(
        block_type=BlockType.SECTION_HEADER,
        bbox=[0.0, 0.1, 0.5, 0.15],
        text="Chapter 1",
        page_idx=0,
        level=1,
        confidence=0.92,
    )
    d = node.to_dict()
    assert d["block_type"] == "section_header"
    assert d["text"] == "Chapter 1"
    assert d["level"] == 1
    assert d["confidence"] == 0.92
    assert d["bbox"] == [0.0, 0.1, 0.5, 0.15]
    # block_id is non-empty
    assert isinstance(d["block_id"], str) and d["block_id"]


def test_from_pages_data_basic():
    pages = {
        0: [
            ([0.0, 0.0, 1.0, 0.1], "INTRODUCTION"),
            ([0.0, 0.1, 1.0, 0.2], "This is a normal paragraph."),
        ],
        1: [
            ([0.0, 0.0, 1.0, 0.1], "CHAPTER 1"),
            ([0.0, 0.1, 1.0, 0.2], "Some body text."),
        ],
    }
    tree = from_pages_data(pages, source_path="doc.pdf")
    assert tree.source_path == "doc.pdf"
    assert len(tree.pages) == 2
    assert tree.pages[0].page_idx == 0
    # All-caps short lines should be classified as SECTION_HEADER
    kinds = {n.block_type for n in tree.pages[0].children}
    assert BlockType.SECTION_HEADER in kinds
    assert BlockType.PARAGRAPH in kinds
    # round-trip through to_dict
    d = tree.to_dict()
    assert "pages" in d and len(d["pages"]) == 2


def test_document_tree_iter_text_blocks():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0, 1, 0.1],
                        text="hello",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_HEADER,
                        bbox=[0, 0, 1, 0.05],
                        text="HEADER",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )
    blocks = tree.iter_text_blocks()
    # Both blocks are yielded by iter_text_blocks; translate_tree is the
    # function that skips headers/footers/numbers (see test below).
    texts = sorted(b.text for b in blocks)
    assert texts == ["HEADER", "hello"]


def test_translate_tree_skips_header_footer_number():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PAGE_HEADER,
                        bbox=[0, 0, 1, 0.05],
                        text="HEADER",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.1, 1, 0.2],
                        text="body",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_FOOTER,
                        bbox=[0, 0.95, 1, 1],
                        text="FOOTER",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_NUMBER,
                        bbox=[0.4, 0.5, 0.6, 0.6],
                        text="42",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )

    async def translator(prompt: str, lang: str) -> str:
        return "t"

    asyncio.run(
        translate_tree(
            tree,
            target_language="French",
            translator=translator,
        )
    )
    # Headers/footers/page-numbers are unchanged
    assert tree.pages[0].children[0].text == "HEADER"
    assert tree.pages[0].children[2].text == "FOOTER"
    assert tree.pages[0].children[3].text == "42"
    # The body paragraph was translated
    assert tree.pages[0].children[1].text == "t"


def test_table_node_to_dict_shape():
    table = TableNode(
        rows=2,
        cols=2,
        page_idx=0,
        bbox=[0, 0, 1, 0.5],
        cells=[
            [
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0, 0, 0.5, 0.25],
                    text="A",
                    page_idx=0,
                ),
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0.5, 0, 1, 0.25],
                    text="B",
                    page_idx=0,
                ),
            ]
        ],
    )
    d = table.to_dict()
    assert d["rows"] == 2 and d["cols"] == 2
    assert d["cells"][0][0]["text"] == "A"


# ---------------------------------------------------------------------------
# Glossary
# ---------------------------------------------------------------------------


def test_glossary_paired_lines_parsing():
    text = "Longer Phrase = Longer Phrase FR\n# this is a comment\nX = Y\n"
    g = Glossary.from_paired_lines(text)
    assert len(g.entries) == 2
    assert g.entries[0].source == "Longer Phrase"
    assert g.entries[0].target == "Longer Phrase FR"
    # Longest-first in prompt block
    block = g.to_prompt_block()
    assert "GLOSSARY" in block
    # The longer entry comes first in the rendered block
    assert block.index("Longer Phrase") < block.index("X -> Y")


def test_glossary_apply_to_text_case_insensitive():
    g = Glossary(entries=[GlossaryEntry(source="Apple", target="Pomme")])
    out = g.apply_to_text("I have an apple. APPLE pie.")
    # Case-insensitive substitution should match both
    assert "Pomme" in out and out.lower().count("pomme") == 2


def test_glossary_apply_to_text_case_sensitive():
    g = Glossary(
        entries=[GlossaryEntry(source="Apple", target="Pomme", case_sensitive=True)]
    )
    out = g.apply_to_text("Apple apple APPLE")
    # Only the exact-case "Apple" gets replaced
    assert out.startswith("Pomme")
    # The other casings remain
    assert "apple" in out and "APPLE" in out


def test_glossary_from_dict_filters_empty():
    g = Glossary.from_dict(
        {
            "entries": [
                {"source": "X", "target": "Y"},
                {"source": "", "target": "Z"},
                {"source": "A", "target": ""},
                "not-a-dict",
            ]
        }
    )
    assert len(g.entries) == 1
    assert g.entries[0].source == "X"


def test_glossary_merge_last_wins():
    a = Glossary(entries=[GlossaryEntry(source="A", target="1")])
    b = Glossary(entries=[GlossaryEntry(source="A", target="2")])
    merged = Glossary.merge([a, b])
    assert any(e.target == "2" for e in merged.entries)


# ---------------------------------------------------------------------------
# Entity memory
# ---------------------------------------------------------------------------


def test_entity_memory_extracts_names_dates_acronyms():
    mem = EntityMemory()
    mem.add_text(
        "Steve Jobs founded Apple on January 9, 2007. The iPhone launched later that year."
    )
    mem.add_text("The NASA team joined IBM and MIT in 2010-03-15.")
    block = mem.to_prompt_block()
    # The dates and acronyms are picked up reliably.
    assert "January 9, 2007" in block
    assert "2010-03-15" in block
    assert "NASA" in block
    assert "IBM" in block
    assert "MIT" in block
    # Proper nouns: at least one of these is present (the regex picks up
    # the parts it can see). "Jobs" is always picked; "Apple" is picked
    # when preceded by whitespace.
    assert "Jobs" in block
    assert "Apple" in block


def test_entity_memory_picks_up_multiword_proper_nouns_with_lead_text():
    # A leading connector word lets the regex see "Steve Jobs" together.
    mem = EntityMemory()
    mem.add_text("We recall that Steve Jobs and Tim Cook worked at Apple Inc.")
    block = mem.to_prompt_block()
    assert "Steve Jobs" in block
    assert "Tim Cook" in block
    assert "Apple" in block


def test_entity_memory_is_empty():
    assert EntityMemory().is_empty()
    mem = EntityMemory()
    mem.add_text("plain text with no entities 1234")
    # 1234 alone isn't a date
    assert mem.is_empty() or len(mem.dates) == 0


def test_entity_memory_merge_combines():
    a = EntityMemory()
    a.add_text("We admire Steve Jobs and Wozniak")
    b = EntityMemory()
    b.add_text("People respect Musk and also Tesla Motors")
    merged = a.merge(b)
    # Both entities end up in the merged set
    assert "Steve Jobs" in merged.names
    assert "Musk" in merged.names
    assert "Wozniak" in merged.names
    # Multi-word names are also captured
    assert "Tesla Motors" in merged.names


# ---------------------------------------------------------------------------
# Handwriting preprocessor
# ---------------------------------------------------------------------------


def test_sauvola_binarize_produces_binary_image():
    img = np.full((100, 100), 200, dtype=np.uint8)
    img[20:40, 20:80] = 30
    out = sauvola_binarize(img, window=15)
    assert set(np.unique(out).tolist()).issubset({0, 255})
    # The dark region should mostly be black in the output
    assert out[25, 50] == 0


def test_estimate_stroke_width_runs():
    # Solid bar, 10px wide
    binary = np.full((50, 50), 255, dtype=np.uint8)
    binary[20:30, 10:40] = 0
    sw = estimate_stroke_width(binary)
    # The distance transform should report a value > 0
    assert sw > 0


def test_normalize_stroke_width_idempotent():
    binary = np.full((50, 50), 255, dtype=np.uint8)
    binary[20:30, 10:40] = 0
    out = normalize_stroke_width(binary, target=4.0)
    # Same shape, still binary
    assert out.shape == binary.shape
    assert set(np.unique(out).tolist()).issubset({0, 255})


def test_is_handwritten_page_dense_text_returns_something():
    # Build a synthetic "handwriting-like" image: low ink density, irregular
    rng = np.random.default_rng(42)
    img = np.full((200, 200), 255, dtype=np.uint8)
    # Add some random sparse dark pixels
    coords = rng.integers(0, 200, size=(200, 2))
    for x, y in coords:
        img[y, x] = 0
    b64 = _arr_to_b64(img)
    # We only care that this doesn't crash; result may be True or False.
    result = is_handwritten_page(b64)
    assert isinstance(result, bool)


def test_sauvola_binarize_matches_pre_hoist_formulation():
    """§1.6 regression: hoisting the ``astype(np.float32)`` cast is semantics-preserving.

    The optimization rebinds ``gray_f32 = gray.astype(np.float32)`` once and
    reuses it across the mean / sqmean / threshold-comparison sites, eliminating
    two redundant float-buffer allocations per page. Assert byte-for-byte
    equivalence with the pre-hoist formulation on a deterministic input so a
    future change to the cast site cannot silently alter the threshold.
    """
    import cv2

    from omniscribe.core.handwriting_preprocessor import sauvola_binarize

    rng = np.random.default_rng(2026)
    gray = rng.integers(0, 256, size=(64, 64), dtype=np.uint8)
    window = 15
    k = 0.2
    r = 128.0

    # Pre-hoist formulation (three separate astype calls).
    mean_old = cv2.boxFilter(gray.astype(np.float32), ddepth=-1, ksize=(window, window))
    sqmean_old = cv2.boxFilter(
        (gray.astype(np.float32)) ** 2, ddepth=-1, ksize=(window, window)
    )
    var_old = np.maximum(sqmean_old - mean_old * mean_old, 0.0)
    std_old = np.sqrt(var_old)
    threshold_old = mean_old * (1.0 + k * (std_old / r - 1.0))
    expected = np.where(gray.astype(np.float32) < threshold_old, 0.0, 255.0).astype(
        np.uint8
    )

    # Hoisted formulation (one astype call, reused).
    actual = sauvola_binarize(gray, window=window, k=k, r=r)

    assert actual.shape == expected.shape
    assert actual.dtype == np.uint8
    # The two formulations must be byte-identical on this deterministic input.
    assert np.array_equal(actual, expected)


def test_handwriting_options_is_noop():
    # An instance with every transformation flag disabled is a no-op.
    assert HandwritingOptions(
        enabled=False,
        binarize=False,
        normalize_stroke_width=False,
        normalize_slant=False,
    ).is_noop()
    # Any flag set means we will do work.
    assert not HandwritingOptions(enabled=True).is_noop()
    assert not HandwritingOptions(binarize=True).is_noop()


def _arr_to_b64(arr: np.ndarray) -> str:
    import base64

    import cv2

    ok, buf = cv2.imencode(".png", arr)
    assert ok
    return base64.b64encode(buf.tobytes()).decode("ascii")


# ---------------------------------------------------------------------------
# Translation tree
# ---------------------------------------------------------------------------


def test_build_context_block_orders_sections():
    g = Glossary(entries=[GlossaryEntry(source="X", target="Y")])
    m = EntityMemory()
    m.add_text("Steve Jobs")
    block = build_context_block(g, m, "the quick brown fox")
    # All three sections appear
    assert "GLOSSARY" in block
    assert "PROPER NOUNS" in block
    assert "PREVIOUS CONTEXT" in block
    # Glossary comes first
    assert block.index("GLOSSARY") < block.index("PROPER NOUNS")
    assert block.index("PROPER NOUNS") < block.index("PREVIOUS CONTEXT")


def test_translate_tree_writes_back_and_preserves_structure():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.SECTION_HEADER,
                        bbox=[0, 0, 1, 0.1],
                        text="Hello",
                        page_idx=0,
                        level=1,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.1, 1, 0.2],
                        text="World",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PAGE_HEADER,
                        bbox=[0, 0.95, 1, 1],
                        text="pg 1",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )

    async def translator(prompt: str, lang: str) -> str:
        return f"[{lang}] {prompt.split('SOURCE:')[-1].strip().splitlines()[0]}"

    asyncio.run(
        translate_tree(
            tree,
            target_language="French",
            translator=translator,
        )
    )
    # The section header and paragraph were translated
    assert tree.pages[0].children[0].text.startswith("[French] Hello")
    assert tree.pages[0].children[1].text.startswith("[French] World")
    # The page header was skipped
    assert tree.pages[0].children[2].text == "pg 1"
    # Translation metadata recorded
    assert "translation" in tree.pages[0].children[0].metadata


def test_translate_tree_sliding_window_propagates():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0, 1, 0.1],
                        text="alpha bravo charlie",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.1, 1, 0.2],
                        text="delta echo foxtrot",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )
    seen: list[str] = []

    async def translator(prompt: str, lang: str) -> str:
        seen.append(prompt)
        # Echo back a long string so the sliding window picks it up
        return ("ok " * 50).strip()

    asyncio.run(
        translate_tree(
            tree,
            target_language="Spanish",
            translator=translator,
            sliding_window_words=10,
        )
    )
    # The second prompt should contain the PREVIOUS CONTEXT section
    assert "PREVIOUS CONTEXT" in seen[1]


def test_translate_tree_dual_translate_chooses_secondary_when_closer():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0, 1, 0.1],
                        text="hi",  # very short source
                        page_idx=0,
                    )
                ],
            )
        ]
    )

    async def primary(prompt: str, lang: str) -> str:
        return "this is a much longer and hallucinated translation that drops nothing"

    async def secondary(prompt: str, lang: str) -> str:
        return "hi-traduit"  # very close in length to "hi"

    asyncio.run(
        translate_tree(
            tree,
            target_language="French",
            translator=primary,
            second_translator=secondary,
            dual_translate=True,
        )
    )
    # The shorter, closer-length secondary should win
    assert tree.pages[0].children[0].text == "hi-traduit"


async def test_translate_node_includes_glossary_and_memory(monkeypatch):
    """When the new optional state fields are populated, they must end up in the prompt."""
    from omniscribe.core import translation as translation_mod

    captured: dict[str, object] = {}

    async def fake_call_llm(**kwargs):
        captured["messages"] = kwargs.get("messages")
        return "translated"

    monkeypatch.setattr(translation_mod, "call_llm", fake_call_llm)

    state = {
        "source_chunk": "Bonjour le monde",
        "target_language": "English",
        "glossary_prompt_block": "GLOSSARY: Bonjour = Hello",
        "entity_memory_prompt_block": "NAMES: Paris",
        "sliding_window": "previously translated text",
    }
    out = await translation_mod.translate_node(state)
    assert out["translated_chunk"] == "translated", out
    messages = captured.get("messages")
    assert isinstance(messages, list) and messages and isinstance(messages[0], dict)
    prompt = messages[0]["content"]
    assert "GLOSSARY: Bonjour = Hello" in prompt
    assert "NAMES: Paris" in prompt
    assert "PREVIOUS CONTEXT" in prompt
    assert "previously translated text" in prompt
    assert "SOURCE TEXT" in prompt


# ---------------------------------------------------------------------------
# HTML writer
# ---------------------------------------------------------------------------


def test_html_writer_emits_semantic_tags():
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.SECTION_HEADER,
                        bbox=[0, 0, 1, 0.1],
                        text="Chapter 1",
                        page_idx=0,
                        level=1,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.1, 1, 0.2],
                        text="Body text.",
                        page_idx=0,
                        confidence=0.85,
                    ),
                    BlockNode(
                        block_type=BlockType.CODE,
                        bbox=[0, 0.2, 1, 0.3],
                        text="print(hi)",
                        page_idx=0,
                    ),
                ],
            ),
            PageTree(
                page_idx=1,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0, 1, 0.1],
                        text="Page 2",
                        page_idx=1,
                    ),
                ],
            ),
        ]
    )
    html = render_html(tree)
    assert "<h1" in html and "Chapter 1" in html
    assert "<p" in html and "Body text." in html
    assert "<pre" in html and "<code>print(hi)</code>" in html
    assert "<!-- PageBreak -->" in html
    assert "data-block-id" in html
    assert 'data-confidence="0.850"' in html
    assert '<section data-page-idx="0">' in html
    assert '<section data-page-idx="1">' in html


def test_html_writer_handles_figure_and_table():
    table = TableNode(
        rows=2,
        cols=2,
        page_idx=0,
        bbox=[0, 0, 1, 0.5],
        cells=[
            [
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0, 0, 0.5, 0.25],
                    text="A",
                    page_idx=0,
                ),
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0.5, 0, 1, 0.25],
                    text="B",
                    page_idx=0,
                ),
            ],
            [
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0, 0.25, 0.5, 0.5],
                    text="C",
                    page_idx=0,
                ),
                BlockNode(
                    block_type=BlockType.PARAGRAPH,
                    bbox=[0.5, 0.25, 1, 0.5],
                    text="D",
                    page_idx=0,
                ),
            ],
        ],
    )
    # Put the table on the tree's `tables` list; the figure goes in page children.
    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.FIGURE,
                        bbox=[0, 0, 1, 0.5],
                        text="caption text",
                        page_idx=0,
                    ),
                ],
            )
        ],
        tables=[table],
    )
    html = render_html(tree)
    assert "<figure" in html and "<figcaption>caption text</figcaption>" in html
    # The tables list is also rendered (figure + table block).
    assert "<table" in html
    # Header row uses <th>, data rows use <td>; the <th> contains a data-block-id attr
    assert ">A</th>" in html and ">C</td>" in html


# ---------------------------------------------------------------------------
# DOCX tree writer
# ---------------------------------------------------------------------------


def test_docx_tree_writer_creates_real_structure():
    from docx import Document

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.SECTION_HEADER,
                        bbox=[0, 0, 1, 0.1],
                        text="Title",
                        page_idx=0,
                        level=1,
                    ),
                    BlockNode(
                        block_type=BlockType.LIST_ITEM,
                        bbox=[0, 0.1, 1, 0.2],
                        text="item",
                        page_idx=0,
                        level=0,
                    ),
                    BlockNode(
                        block_type=BlockType.CODE,
                        bbox=[0, 0.2, 1, 0.3],
                        text="x = 1",
                        page_idx=0,
                    ),
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0.3, 1, 0.4],
                        text="body",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )
    stream = convert_tree_to_docx(tree)
    doc = Document(stream)
    # Heading 1 should be present
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    assert "Title" in headings
    # Code line preserved
    assert any("x = 1" in p.text for p in doc.paragraphs)
    # Body text present
    assert any("body" in p.text for p in doc.paragraphs)
    # List item present
    assert any("item" in p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Tree JSON export
# ---------------------------------------------------------------------------


def test_tree_export_json_round_trip():
    import json as _json

    tree = DocumentTree(
        pages=[
            PageTree(
                page_idx=0,
                children=[
                    BlockNode(
                        block_type=BlockType.PARAGRAPH,
                        bbox=[0, 0, 1, 0.1],
                        text="hi",
                        page_idx=0,
                    ),
                ],
            )
        ]
    )
    raw = export_json(tree)
    data = _json.loads(raw)
    assert "pages" in data
    assert data["pages"][0]["children"][0]["text"] == "hi"
    # Bytes variant
    b = export_json_bytes(tree)
    assert isinstance(b, bytes)
    assert _json.loads(b.decode("utf-8"))["pages"][0]["page_idx"] == 0


# ---------------------------------------------------------------------------
# Dual translator
# ---------------------------------------------------------------------------


def test_dual_translate_chooser():
    async def primary(prompt: str, lang: str) -> str:
        return "this is a very long and unhelpful translation that pads the output"

    async def secondary(prompt: str, lang: str) -> str:
        return "salut mon ami"

    def build_prompt(text: str, lang: str) -> str:
        return f"Translate to {lang}: {text}"

    chosen, meta = asyncio.run(
        dual_translate(
            "hi",
            target_language="French",
            primary=primary,
            secondary=secondary,
            build_prompt=build_prompt,
        )
    )
    assert meta["strategy"] == "dual"
    # secondary (3 words, very close to "hi") should be chosen
    assert chosen == "salut mon ami"
    assert meta["primary_length_ratio"] > meta["secondary_length_ratio"]


def test_dual_translate_falls_back_when_no_prompt_builder():
    async def primary(prompt: str, lang: str) -> str:
        return "primary result"

    chosen, meta = asyncio.run(
        dual_translate(
            "hi", target_language="French", primary=primary, secondary=primary
        )
    )
    assert chosen == "primary result"
    assert meta["strategy"] == "single"


# ---------------------------------------------------------------------------
# TrOCR heuristic confidence
# ---------------------------------------------------------------------------


def test_trocr_heuristic_confidence():
    assert _heuristic_confidence("") == 0.0
    assert _heuristic_confidence("bcdfg") == 0.2  # no vowel
    # Two words sits in the 0.7 band
    assert _heuristic_confidence("hello world") == 0.7
    # Three or more words triggers the higher confidence band
    assert _heuristic_confidence("one two three four") == 0.85


def test_trocr_engine_is_available():
    from omniscribe.core.trocr_engine import TrOCREngine

    eng = TrOCREngine()
    # The function must return a bool; the actual value depends on env
    assert isinstance(eng.is_available(), bool)


# ---------------------------------------------------------------------------
# NLLB language code resolver
# ---------------------------------------------------------------------------


def test_nllb_resolve_nllb_code_known():
    assert resolve_nllb_code("French") == "fra_Latn"
    assert resolve_nllb_code("english") == "eng_Latn"
    assert resolve_nllb_code("Chinese") == "zho_Hans"
    # Already a code
    assert resolve_nllb_code("deu_Latn") == "deu_Latn"
    # Unknown -> default to English
    assert resolve_nllb_code("Klingon") == "eng_Latn"


def test_nllb_language_code_map_has_basics():
    for name in ("english", "spanish", "french", "german", "chinese", "japanese"):
        assert name in LANGUAGE_CODE_MAP
        assert "_" in LANGUAGE_CODE_MAP[name]
