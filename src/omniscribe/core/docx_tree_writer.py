"""DOCX export for ``DocumentTree``.

This is the structured counterpart to :func:`convert_markdown_to_docx` — the
markdown path stays for backward compat. The tree path walks a
:class:`DocumentTree` and emits real Word objects: heading styles, native
tables, embedded images, code blocks, definition lists.
"""

from __future__ import annotations

import contextlib
import io
from typing import TYPE_CHECKING, Any, cast

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from omniscribe.core.block_tree import BlockNode, DocumentTree, TableNode


def convert_tree_to_docx(tree: DocumentTree) -> io.BytesIO:
    """Render a :class:`DocumentTree` to a Word document and return a stream."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for page in tree.pages:
        for node in page.children:
            _render_block(doc, node)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _render_block(doc: Any, node: BlockNode) -> None:
    bt = node.block_type.value
    if bt == "section_header":
        level = max(1, min(6, node.level or 1))
        h = doc.add_heading(node.text, level=level)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
    elif bt == "list_item":
        style_name = "List Bullet" if node.level == 0 else "List Bullet 2"
        p = doc.add_paragraph(node.text, style=style_name)
        p.paragraph_format.space_after = Pt(3)
    elif bt == "code":
        p = doc.add_paragraph()
        run = p.add_run(node.text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    elif bt == "equation":
        p = doc.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run("$ " + node.text + " $")
        run.italic = True
    elif bt == "figure":
        if node.metadata.get("image_bytes"):
            p = doc.add_paragraph()
            run = p.add_run()
            with contextlib.suppress(Exception):
                run.add_picture(
                    io.BytesIO(node.metadata["image_bytes"]), width=Inches(5.5)
                )
        if node.text:
            cap = doc.add_paragraph(node.text)
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    elif bt == "table":
        _render_table(doc, cast("TableNode", node))
    elif bt in ("page_header", "page_footer", "page_number"):
        # Don't render these as body content.
        return
    elif bt == "key_value":
        p = doc.add_paragraph()
        run = p.add_run(node.metadata.get("key", "") + ": ")
        run.bold = True
        p.add_run(node.text)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if node.spans:
            for sp in node.spans:
                run = p.add_run(sp.text)
                if sp.bold:
                    run.bold = True
                if sp.italic:
                    run.italic = True
                if sp.code:
                    run.font.name = "Courier New"
        else:
            p.add_run(node.text)


def _render_table(doc: Any, table_node: TableNode) -> None:
    rows = max(1, table_node.rows)
    cols = max(1, table_node.cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(table_node.cells):
        for c_idx, cell_node in enumerate(row):
            if r_idx >= rows or c_idx >= cols:
                continue
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_node.text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
