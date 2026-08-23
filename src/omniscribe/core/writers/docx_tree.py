"""DOCX export for ``DocumentTree``.

This is the structured counterpart to :func:`convert_markdown_to_docx` — the
markdown path stays for backward compat. The tree path walks a
:class:`DocumentTree` and emits real Word objects: heading styles, native
tables, embedded images, code blocks, definition lists.
"""

from __future__ import annotations

import contextlib
import io
from typing import TYPE_CHECKING, Any, cast

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from omniscribe.core.writers.exporter_base import BaseDocumentExporter

if TYPE_CHECKING:
    from omniscribe.core.block_tree import BlockNode, DocumentTree, TableNode


class DocxTreeExporter(BaseDocumentExporter):
    """Document exporter producing structured Word (.docx) documents from DocumentTree."""

    def export_tree(self, tree: DocumentTree, **kwargs: Any) -> io.BytesIO:
        """Render a DocumentTree to a Word document stream."""
        return convert_tree_to_docx(tree)


def convert_tree_to_docx(tree: DocumentTree) -> io.BytesIO:
    """Render a :class:`DocumentTree` to a Word document and return a stream."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    rendered_tables: set[str | int] = set()
    for page in tree.pages:
        for node in page.children:
            if hasattr(node, "rows") and hasattr(node, "cells"):
                rendered_tables.add(id(node))
                if getattr(node, "block_id", None):
                    rendered_tables.add(node.block_id)
            _render_block(doc, node, rendered_tables)
    for table in tree.tables:
        if id(table) not in rendered_tables and (
            not table.block_id or table.block_id not in rendered_tables
        ):
            _render_table(doc, table)
            rendered_tables.add(id(table))
            if table.block_id:
                rendered_tables.add(table.block_id)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _render_block(
    doc: Any,
    node: BlockNode | TableNode | Any,
    rendered_tables: set[str | int] | None = None,
) -> None:
    if hasattr(node, "rows") and hasattr(node, "cells"):
        if rendered_tables is not None:
            rendered_tables.add(id(node))
            if getattr(node, "block_id", None):
                rendered_tables.add(node.block_id)
        _render_table(doc, cast("TableNode", node))
        return
    if not hasattr(node, "block_type"):
        return
    bt = (
        node.block_type.value
        if hasattr(node.block_type, "value")
        else str(node.block_type)
    )
    if bt == "section_header":
        level = max(1, min(6, getattr(node, "level", 1) or 1))
        h = doc.add_heading(node.text, level=level)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
    elif bt == "list_item":
        style_name = (
            "List Bullet" if getattr(node, "level", 0) == 0 else "List Bullet 2"
        )
        p = doc.add_paragraph(node.text, style=style_name)
        p.paragraph_format.space_after = Pt(3)
    elif bt == "code":
        p = doc.add_paragraph()
        run = p.add_run(node.text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    elif bt == "equation":
        p = doc.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run("$ " + node.text + " $")
        run.italic = True
    elif bt == "figure":
        metadata = getattr(node, "metadata", {}) or {}
        if metadata.get("image_bytes"):
            p = doc.add_paragraph()
            run = p.add_run()
            with contextlib.suppress(Exception):
                run.add_picture(io.BytesIO(metadata["image_bytes"]), width=Inches(5.5))
        if getattr(node, "text", ""):
            cap = doc.add_paragraph(node.text)
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    elif bt == "table":
        if rendered_tables is not None:
            rendered_tables.add(id(node))
            if getattr(node, "block_id", None):
                rendered_tables.add(node.block_id)
        if hasattr(node, "cells") and getattr(node, "cells", None):
            _render_table(doc, cast("TableNode", node))
        elif hasattr(node, "text") and node.text:
            p = doc.add_paragraph(node.text)
            p.paragraph_format.space_after = Pt(6)
    elif bt in ("page_header", "page_footer", "page_number"):
        # Don't render these as body content.
        return
    elif bt == "key_value":
        metadata = getattr(node, "metadata", {}) or {}
        p = doc.add_paragraph()
        run = p.add_run(metadata.get("key", "") + ": ")
        run.bold = True
        p.add_run(getattr(node, "text", ""))
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        spans = getattr(node, "spans", None)
        if spans:
            for sp in spans:
                run = p.add_run(sp.text)
                if sp.bold:
                    run.bold = True
                if sp.italic:
                    run.italic = True
                if sp.code:
                    run.font.name = "Courier New"
        else:
            p.add_run(getattr(node, "text", ""))


def _render_table(doc: Any, table_node: TableNode | BlockNode | Any) -> None:
    cells = getattr(table_node, "cells", None)
    if not cells or not isinstance(cells, (list, tuple)):
        text = getattr(table_node, "text", "")
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        return

    num_rows = getattr(table_node, "rows", None)
    num_cols = getattr(table_node, "cols", None)
    if num_rows is None or num_rows <= 0:
        num_rows = len(cells)
    if num_cols is None or num_cols <= 0:
        num_cols = max(
            (len(r) for r in cells if isinstance(r, (list, tuple))), default=0
        )

    if num_rows <= 0 or num_cols <= 0:
        text = getattr(table_node, "text", "")
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        return

    rows = max(1, num_rows)
    cols = max(1, num_cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(cells):
        if not isinstance(row, (list, tuple)):
            continue
        for c_idx, cell_node in enumerate(row):
            if r_idx >= rows or c_idx >= cols:
                continue
            cell = table.cell(r_idx, c_idx)
            cell_text = getattr(
                cell_node, "text", str(cell_node) if cell_node is not None else ""
            )
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


__all__ = [
    "DocxTreeExporter",
    "convert_tree_to_docx",
]
